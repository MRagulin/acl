import os, shutil, sys
from pathlib import Path, PurePosixPath
from django.core.files.storage import FileSystemStorage
from django.apps import apps
from django.contrib import messages
import re
from django.conf import settings
from django.db.utils import IntegrityError, DataError
import datetime
import time
import uuid
import ipaddress
from django.views import View
from django.http import JsonResponse, HttpResponseRedirect
import logging
import xlrd
import tempfile
from django.shortcuts import reverse, redirect
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
import socket
import codecs
import json
from time import sleep
import git
from shutil import copyfile
from django.core.cache import cache
import random, hashlib, datetime
COMMIT_MESSAGE = '[ACL PORTAL] Add acl-MD file'


FUN_SPEED = 0
BASE_DIR = Path(__file__).resolve().parent.parent
LOCAL_UID = None

FORM_APPLICATION_KEYS = ['acl_create_info.html', 'acl_internal_resources.html', 'acl_dmz_resources.html', 'acl_external_resources.html', 'acl_traffic.html', 'acl_approve.html']
FORM_URLS = ["acldemo_urls", "aclcreate_urls", "aclinternal_urls", "acldmz_urls", "aclexternal_urls", "acltraffic_urls", "acloverview_urls", 'acl_approve_urls', "acl_pending_urls"]
POST_FORM_KEYS = ['name', 'email', 'tel', 'department', 'project', 'link', 'd_form', 'd_start', 'd_complate']
POST_FORM_EMPTY = ['on', '', None]
JSON_DUMPS_PARAMS = {
    'ensure_ascii': False
}
left_rule = {'<': ':', '^': ':', '>': '-'}
right_rule = {'<': '-', '^': ':', '>': ':'}

contact_column = ["Параметр", "Значение"]
contact_table = ["ФИО", "E-mail", "Телефон", "Отдел/Управление", "Информационная система", "Описание/архитектура проекта", "Дата заполнения", "Дата ввода в эксплуатацию", "Дата вывода из эксплуатации"]
standart_column = ["IP-адрес", "Маска подсети/Префикс", "Описание"]
traffic_column = ["Hostname (Источник)", "IP Address (Источник)	", "Hostname (Назначение)", "IP Address (Назначение)", "Protocol/Port (Назначение)", "Описание (цель)"]


logger = logging.getLogger(__name__)


class BaseView(View):
    def dispatch(self, request, *args, **kwargs):
        try:
            response = super().dispatch(request, *args, **kwargs)
        except Exception as e:
            logger.error('{}|{}|{}|{}'.format(request.path, str(e), request.META.get('REMOTE_ADDR'), datetime.datetime.today().strftime('%Y-%m-%d-%H:%M:%S')))
            messages.error(request, str(e))
            return HttpResponseRedirect(reverse('acldemo_urls')) #self.__response({'errorMessage': str(e)}, status=400)

        if isinstance(response, (dict, list)):
            return self.__response(response)
        else:
             return response

    @staticmethod
    def __response(data, *, status=200):
        return JsonResponse(
            data,
            status=status,
            safe=not isinstance(data, list),
            json_dumps_params=JSON_DUMPS_PARAMS
        )

def isvalidip(ip)-> bool:
        l = len(str(ip))
        if (l == 0) or (l > 15): return False
        s = str(ip).split('.')
        if len(s) >= 3:
            return True
        else:
            return False


def get_client_ip(request)->str:
    """Получение IP адреса клиента"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip


def ip_status(ip=None)->dict:
    """Проверка типа IP адресса"""
    data = {}
    data['ip'] = False
    try:
        ip = ipaddress.ip_address(ip)
    except ValueError:
        return data
    data['ip'] = True

    if ip.is_reserved:
        data['type'] = 3
        return data

    if ip.is_loopback:
        data['type'] = 4
        return data
    if ip.is_multicast:
        data['type'] = 5
        return data

    if ip.is_global:
        data['type'] = 1
    elif ip.is_private:
        data['type'] = 2
    else:
        data['type'] = 0
    return data


def request_handler(request, namespace=''):
    """Функция для заполнения глобального массива LOCAL_STORAGE из POST параметров файлов acl*"""
    INFINITY = 'Нет'
    LOCAL_STORAGE = {}
    cnt_key = 0
    empty_key = 0
    if namespace == FORM_APPLICATION_KEYS[0]: #first
            LOCAL_STORAGE[namespace] = []
            for idx, post_key in enumerate(POST_FORM_KEYS):
                if idx == len(POST_FORM_KEYS) - 1:
                    if request.POST.get(post_key) in POST_FORM_EMPTY:
                        LOCAL_STORAGE[namespace].append(INFINITY)
                        continue

                if request.POST.get(post_key) not in POST_FORM_EMPTY:
                     LOCAL_STORAGE[namespace].append(request.POST.get(post_key))
                else:
                    LOCAL_STORAGE[namespace].append(INFINITY)
                   # return False
            if namespace == FORM_APPLICATION_KEYS[0]:
                    if request.POST.get('action_make_docx') == 'on':
                        request.session['ACT_MAKE_DOCX'] = True
                    else:
                        if 'ACT_MAKE_DOCX' in request.session:
                            del request.session['ACT_MAKE_DOCX']

                    if request.POST.get('action_make_git') == 'on':
                        request.session['ACT_MAKE_GIT'] = True
                    else:
                        if 'ACT_MAKE_GIT' in request.session:
                            del request.session['ACT_MAKE_GIT']
                    request.session.modified = True

    else:
        if namespace == FORM_APPLICATION_KEYS[-1]: #last
            str_pattern = 'input__domain_source'
        else:
            str_pattern = 'input__ip'

        for k, v in request.POST.items():
                if 'input_' in str(k):
                        if str_pattern in str(k):
                            if namespace in LOCAL_STORAGE:
                                LOCAL_STORAGE[namespace].append([v])
                                cnt_key += 1
                            else:
                                LOCAL_STORAGE[namespace] = [[v]]
                        else:
                            if v != '':
                                LOCAL_STORAGE[namespace][cnt_key].append(v)
                            else:
                                if k == 'd_complate':
                                    LOCAL_STORAGE[namespace][cnt_key].append(INFINITY)
                                else:
                                    empty_key += 1

        # if empty_key >= 2:
        #         del LOCAL_STORAGE[namespace]
                                #return False
    return LOCAL_STORAGE


def IP2Int(ip):
    """Function convert IP to integer"""
    o = list(map(int, ip.split('.')))
    res = (16777216 * o[0]) + (65536 * o[1]) + (256 * o[2]) + o[3]
    return res


def upload_file_handler(request, functionhandler = None):
    """Функция обработки загрузки файлов и вызова функции для парсинга xls"""
    result = {}
    if 'input--file--upload' in request.FILES:
        UPLOAD_PATH = tempfile.gettempdir() #os.path.join(BASE_DIR, 'upload')
        myfile = request.FILES['input--file--upload']
        fs = FileSystemStorage(location=UPLOAD_PATH)
        fs.save(myfile.name, myfile)
        uploaded_file_url = os.path.join(UPLOAD_PATH, myfile.name) #bug with persone encode
        print('Upload file to: {}'.format(uploaded_file_url))
    else:
        result['error'] = "There is error upload file"
        return result
    if uploaded_file_url == '':
       return {'error': 'There is error upload file'}

    if 'ext.' in uploaded_file_url or \
       'aktur' in uploaded_file_url or \
        'alfatrah.ru' in uploaded_file_url:
        functionhandler = ExtractDataDns

    if functionhandler is not None:
        result = functionhandler(uploaded_file_url)
    else:
        result = ExtractDataXls(request, uploaded_file_url).execute_file_parsing()
        # race condition
    time.sleep(1)
    try:
            os.remove(uploaded_file_url)
            if settings.DEBUG:
                print('Удаление файла: {}'.format(uploaded_file_url))
    finally:
            uploaded_file_url = None
    if result > 0:
            return {'ok': 'Добавлено новых значений: {}'.format(result)}
    return {'error': 'Данных для добавления - нету'}


def count_perf(f):
    """Декоратор для измерения скорости поиска"""
    def wraper(*args, **kwargs):
        global FUN_SPEED
        time_init = datetime.datetime.now()
        result = f(*args, **kwargs)
        time_end = datetime.datetime.now()
        total = time_end - time_init
        FUN_SPEED = total
        #print('Time: {}'.format(total))
        return result
    return wraper

@count_perf
def DeepSearch(request, string: str = ''):
    """Функция для анализа типа данных в запросе и поиск по структуре"""
    result, tmp = '', string
    Iplist = apps.get_model('ownerlist', 'Iplist')

    if re.match(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", tmp):
        result = Iplist.objects.filter(ipv4=tmp)

    if not result:
            if re.match(r"^\d{1,3}\,\d{1,3}\,\d{1,3}\,\d{1,3}$", tmp):
                tmp = string.replace(',', '.')
            if tmp:
                result = Iplist.objects.filter(ipv4__contains=tmp)[:5]
            # else:
            #     result = Iplist.objects.filter(ipv4__contains=string)[:5]
            if result:
                messages.add_message(request, messages.INFO,
                                     'По запросу {} ничего не найдено, но мы нашли похожую информацию:'.format(string))
    if not result:
        if re.match(r"[a-zA-Z0-9][a-zA-Z0-9-._]{1,61}", tmp):

            result = Iplist.objects.filter(hostname=tmp)[:5]

            if not result:
                if '://' in tmp:
                    tmp = tmp.split('://')[1]

                if ('.vesta.ru' in tmp or '.alfastrah.ru' in tmp or '.dyn.vesta.ru' in tmp):
                    tmp = tmp.split('.')[0]


                result = Iplist.objects.filter(hostname__icontains=tmp)[:5]
                if result:
                    messages.add_message(request, messages.INFO,
                                         'По запросу {} ничего не найдено, но мы нашли похожую информацию:'.format(
                                             string))
                if not result:
                    result = Iplist.objects.filter(comment__icontains=tmp)[:5]
                    if result:
                        messages.add_message(request, messages.INFO,
                                             'По запросу {} ничего не найдено, но мы нашли похожую информацию:'.format(
                                                 string))
                    else:
                        try:
                            tmp = socket.gethostbyname(tmp)
                        except:
                            pass
                        if tmp:
                                result = Iplist.objects.filter(ipv4__contains=tmp)[:5]
                        if result:
                                messages.add_message(request, messages.INFO,
                                                     'По запросу {} ничего не найдено, но мы нашли похожую информацию:'.format(
                                                         string))




        #result = Iplist.objects.filter(ipv4__contains=string)[:5]
    #time.sleep(5)
    return result



def write_history(request, string, status) -> None:
    """Сохранять историю поиска, для улучшения качества поиска"""
    hc = apps.get_model('ownerlist', 'HistoryCall')
    ip = apps.get_model('ownerlist', 'Iplist')

    ip_object, obj = ip.objects.get_or_create(ipv4=request.META.get('REMOTE_ADDR')) #IP-адрес пользователя
    hc_object = hc.objects.create(string=string,
                                         ipv4=ip_object,
                                         username= request.user,
                                         status=status)
    return hc_object


def search_text(request=None, string: str = '') -> dict:
    """ Функция для поиска данных в БД"""
    global FUN_SPEED
    result = DeepSearch(request, string)
    context = {'SearchFor': string}
    context['Data'] = result
    context['TakeTime'] = FUN_SPEED
    context['Info'] = ''
    FUN_SPEED = 0
    write_history(request, string, bool(result))
    return context


class ExtractDataXls:
    """Основной класс для анализа xls файла"""
    def __init__(self, request=None, filename=''):
        self.ip_addr_idx = 1
        self.count_total: int = 0 #total records in db
        self.error_count: int = 0 #total errors
        self.rb = xlrd.open_workbook(filename, formatting_info=True, encoding_override='utf-8')

        self.current_page = None
        self.sheet_tags = self.rb.sheet_names()
        self.Vlans = apps.get_model('ownerlist', 'Vlans')
        self.Tags = apps.get_model('ownerlist', 'Tags')
        self.Iplist = apps.get_model('ownerlist', 'Iplist')
        self.Owners = apps.get_model('ownerlist', 'Owners')
        self.page_headers = ['ответственный', 'комменты', 'ip address', 'Имя сервера', 'отвеcтвенный', 'nat inside']
        self.fio_exclude_list = ['гусев', 'оксенюк', 'северцев', 'егоров', 'совинский', 'огнивцев', 'допиро', 'мюлекер', 'уволен', 'иренов', 'казаков', 'куслеев']



    def execute_file_parsing(self):
        """Выбираем парсер на основе имени страницы"""
        result = 0
        for self.sheet_tag in self.sheet_tags:
            self.current_page = self.rb.sheet_by_name(self.sheet_tag)
            if self.current_page.nrows == 0:  # Count row
                if settings.DEBUG:
                    print("Страница {} пустая, пропущено...".format(self.current_page))
                    return 0
            # else:
            #     if settings.DEBUG:
            #         print('Извлекаем данные из: {} ...'.format(self.sheet_tag))
            #         print('*'*60)
            if self.sheet_tag == 'VLAN DESCRIPTION':
                    result += self.ExtractVlanInfo()
            elif self.sheet_tag == 'VLAN_CORE_ACI':
                    result += self.ExtractVlanInfo(name_idx=6, location_idx=3, vlan_idx=2, subnet_idx=4, mask_idx=5, tag1_idx=7, tag2_idx=8)
            elif self.sheet_tag == '10.255.10.0 (NGNX-Serv)':
                result += self.ExtractIPInfo(domain_idx=0, ip_idx=1, owner_idx=2, comment_idx=3)
            elif self.sheet_tag == '213.33.175.0 _24':
                result += self.ExtractIPInfo(domain_idx=1, ip_idx=2, owner_idx=4, comment_idx=5)
            elif self.sheet_tag == 'активка 172.16.82.X':
                result += self.ExtractIPInfo(domain_idx=1, ip_idx=0, owner_idx=2, comment_idx=3)
            elif self.sheet_tag == '195.239.64.хх':
                result += self.ExtractIPInfo(domain_idx=0, ip_idx=1, owner_idx=3, comment_idx=4)
            elif self.current_page.ncols == 4:
                    result += self.ExtractIPInfo()
            else:
                if settings.DEBUG:
                        print("Страница содержит другое количество колонок <> 4, {} анализируем...".format(self.current_page.ncols))
                result += self.PageStructAnalyzer(self.current_page)
        return result


    def is_row_empty(self, row) -> bool:
        """ Проверяем пустая ли запись"""
        result = True
        for d in row:
            if d != '':
                result = False
                break
        return result

    def get_ip_from_page(self, page) -> str:
        """Получаем полный IP из имени страницы"""
        try:
            ip = re.findall(r"(\d{1,3})", page)
            return ".".join(ip)
        except:
            return ""



    def ExtractVlanInfo(self, name_idx=1, location_idx=2, vlan_idx=3, subnet_idx=4, mask_idx=5, tag1_idx=6, tag2_idx=7)->int:
        """Парсер страницы с описанием VLAN"""
        row_index: int = 0
        tags: list = []
        for row_idx in range(self.current_page.nrows):
                        row = self.current_page.row_values(row_idx)
                        if row_idx == 0 or self.is_row_empty(row):
                            continue

                        if type(row[vlan_idx]) == float:
                            vlan = int(round(row[vlan_idx]))
                        elif type(row[vlan_idx]) == str:
                             try:
                                   vlan = int(round(float(row[vlan_idx])))
                             except ValueError:
                                    vlan = 0

                        if str(row[subnet_idx]).find('/') > 0:
                                subnet = str(row[subnet_idx]).split('/')
                                subnet, mask = subnet[0], int(subnet[1])
                        else:
                            try:
                                if len(str(row[subnet_idx])) > 15:
                                    subnet = str(row[subnet_idx]).split('\n')[0] #Bug fig, if a couple value in row
                                else:
                                    subnet = str(row[subnet_idx])
                            except ValueError:
                                subnet = 0


                            try:
                                if len(str(row[mask_idx])) > 4:
                                    mask = str(row[mask_idx]).split('\n')[0] #Bug fig, if a couple value in row
                                    mask = int(round(float(mask)))
                                else:
                                    mask = int(round(float(row[mask_idx]))) or 0
                            except ValueError:
                                mask = 0

                        vlan_info, _ = self.Vlans.objects.get_or_create(
                        name=str(row[name_idx]),
                        location=str(row[location_idx]),
                        vlan=vlan,
                        subnet=subnet,
                        mask=mask,
                        )
                        # if created obj
                        if _:
                                self.count_total += 1

                        try:
                            tags.append(self.sheet_tag)
                            tags.append(row[tag1_idx])
                            tags.append(row[tag2_idx])

                            for tag in tags:
                                if (tag != '') and len(tag) > 1:
                                            if len(str(tag).split('.')) >= 3:  # If tag as Gateway's IP
                                                tag = "Gateway: {}".format(tag)
                                            tag_info, _ = self.Tags.objects.get_or_create(name=str(tag).rstrip())
                                            if tag_info not in vlan_info.tags.all():
                                                vlan_info.tags.add(tag_info)
                                                self.count_total += 1
                        except:
                            pass

                        finally:
                            tags.clear()

        if settings.DEBUG:
            print("Добавленно {} новых записей в БД.".format(self.count_total))

        return self.count_total

    def ExtractIPInfo(self, domain_idx=0, ip_idx=1, owner_idx=2, comment_idx=3, stop_recurse = False, HasTags=[])->int:
        tags: list = []
        self.error_count = 0
        Header_POS = 0
        created = None

        ip_addr = ''
        domain = ''
        owner = ''
        comment = ''



        # if HasTags is not None:
        #     if len(HasTags) > 0:
        #         tags = HasTags

        #print(self.sheet_tag)
        #return 0
        self.count_total = 0
        for row_idx in range(self.current_page.nrows):
            row = self.current_page.row_values(row_idx)
            if not self.is_row_empty(row):
                # if self.sheet_tag == 'Орел':
                #     print ('')
                if len(row) >= 3:
                    if row_idx in range(0, 3):
                        if row[row_idx] in self.page_headers: #Пропускаем заголовки
                            Header_POS = row_idx
                            continue
                if isvalidip(row[ip_idx]):
                    ip_addr = row[ip_idx]
                elif len(str(row[ip_idx])) <= 5: #15.0
                            try:
                                tmp = int(round(float(row[ip_idx]))) or 0
                            except ValueError:
                                tmp = 0
                                continue
                            else:
                                if tmp > 0:
                                    ip_addr = self.get_ip_from_page(str(self.sheet_tag))
                                    if ip_addr != '':
                                            ip_addr = ip_addr + '.' + str(tmp)
                else:
                    self.error_count += 1
                    if self.error_count >= 5:
                        if not stop_recurse:
                            if settings.DEBUG:
                                print("Много ошибок на странице, провёдем анализ страницы ...")
                            self.PageStructAnalyzer(self.current_page)
                        else:
                            if settings.DEBUG:
                                print("********************Ошибка при анализе странице {}****************".format(self.sheet_tag))
                        return 0
                    continue
                try:
                    if domain_idx is not None:
                        domain = row[domain_idx]
                except:
                    domain = ''


                try:
                    if owner_idx is not None:
                        owner = row[owner_idx]
                except:
                    owner = ''

                try:
                    if comment_idx is not None:
                        comment = row[comment_idx]
                except:
                    comment = ''

                if len(HasTags) > 0:
                    for tag in HasTags:
                        try:
                            tmp = row[int(tag)]
                            if tmp != '':
                                name = self.current_page.cell_value(Header_POS, int(tag))
                                if name != '':
                                    tmp = name + ":"+tmp
                                if tmp not in tags:
                                    tags.append(tmp)
                        except:
                            pass
                if self.sheet_tag not in tags:
                    tags.append(self.sheet_tag)

#-----------------------------------------------------------------------------------------------------------------------
                try:
                    if comment and re.match(r"([а-яА-Я\.\s()]){5,}", str(comment)):
                        tmp = comment.lower().strip()
                        exists = list(filter(lambda s: s in tmp, self.fio_exclude_list))
                        #exists = any(substring in string for string in strings)
                        #if comment.lower() in self.fio_exclude_list:
                        if len(exists) > 0:
                            owner, comment = comment, owner
                except:
                    pass
#-----------------------------------------------------------------------------------------------------------------------

                try:
                    if owner:
                        if owner.find('://') != -1:
                            owner, comment = comment, owner
                except:
                    pass

                if ip_addr != '':
                   # if not self.isvalidip(ip_addr):

                       # continue
                    # if comment_idx != -1:
                    #     print("{} {} {} {} ".format(domain, ip_addr, row[owner_idx], row[comment_idx]))
                    # else:
                    if (owner == '') or len(owner) <= 1:
                        owner_info = None
                    else:
                        #try:
                            owner_info, created = self.Owners.objects.get_or_create(username = owner)
                        #except:
                            #owner_info = self.Owners.get_default_owner()

                    try:
                        ip_info, created = self.Iplist.objects.get_or_create(
                            ipv4 = ip_addr,
                            hostname = domain,
                            owner = owner_info,
                            comment = comment
                        )
                    except IntegrityError:
                        ip_info = self.Iplist.objects.get(ipv4=ip_addr)
                        ip_info.ipv4 = ip_addr
                        ip_info.hostname = domain
                        ip_info.owner = owner_info
                        ip_info.comment = comment
                        ip_info.save()

                    except DataError:
                        print("- Ошибка данных: {} на странице: {}".format(ip_addr, self.sheet_tag))
                        continue

                    if created:
                        self.count_total += 1


                    try:
                        for tag in tags:
                            if (tag != '') and len(tag) > 1:
                                tag_info, created = self.Tags.objects.get_or_create(name=str(tag).rstrip())
                                if tag_info not in ip_info.tags.all():
                                    ip_info.tags.add(tag_info)
                    except:
                        pass

                    finally:
                            #print("{} {} {} {}-> {}".format(domain, ip_addr, owner, comment, " [" + ' '.join(tags) + "]"))

                            tags.clear()
        if settings.DEBUG:
            print("Страница: {} записано: {}".format(self.sheet_tag, self.count_total))
        return self.count_total

    def PageStructAnalyzer(self, page, DEBUG = False)-> None:
        """Написанный на коленке анализатор данных на странице в xls"""
        is_domain = 0
        is_ip = 0
        is_owner = 0
        is_comment = 0
        is_add = 0
        col_index ={}
        #RowIndex = 0
        #RowBlackListByNumber = []
        skip_domain = False
        skip_owner = False
        skip_ip = False
        skip_commnet = False

        col_index['domain'] = None
        col_index['ip'] = None
        col_index['owner'] = None
        col_index['comment'] = None

        Tags = []
        #col_index = {}
        # if self.sheet_tag == '172.16.88.Х Avaya':
        #     print('')
        if not (page.ncols > 0):  # and page.ncols < 6
                print("Page {} has been skipped, unknown amount col {}".format(self.sheet_tag, page.ncols))
        else:
                for idx_col in range(page.ncols):
                    col_stat = {'Unknown': 0,
                                'is_domain': 0,
                                'is_ip': 0,
                                'is_owner': 0,
                                'is_comment': 0,
                                'is_tag': 0}

                    #RowIndex = 0
                    if idx_col >=8: #Skip comment in tables
                        continue

                    for index, col in enumerate(page.col_values(idx_col)):
                        if col == '':
                            continue
                        if index in range(0, 2):
                            if col in self.page_headers:
                                continue
                        try:
                            if len(str(col)) <= 1:
                                continue
                            if (re.match(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", str(col)) or
                                re.match(r"^(\d+.\d)|(\d)$", str(col))):
                                    if not skip_ip:
                                        col_stat['is_ip'] += 1
                                        continue
                                    else:
                                      if  idx_col not in Tags:
                                            Tags.append(idx_col)
                                      continue
                        except TypeError:
                            pass


                        try:
                            if (skip_domain == False) and re.match(r"([a-zA-Z0-9\-\.\\()\_]){4,}", str(col)):  # [А-я]+
                                    col_stat['is_domain'] += 1
                                    continue
                        except TypeError:
                                pass


                        try:
                            if (skip_owner == False) and re.match(r"([а-яА-Я\.\s()]){5,}", str(col)):
                                   col_stat['is_owner'] += 1
                                   continue
                        except:
                            pass

                                #else:
                        if len(str(col)) > 3:
                                        t = str(col).lower().find('vlan')
                                        if t > 0: pass
                                        #     str_sum = len(col)
                                        #     #col_stat['is_tag'] += 1
                                        #     try:
                                        #         #cut_vlan = col.split(' ')
                                        #         if idx_col not in Tags:
                                        #             Tags.append(idx_col)
                                        #             continue
                                        #     except:
                                        #         pass
                                            # in ['vlan', 'vlan name', 'location', ]
                                        else:
                                            if not skip_commnet:
                                                col_stat['is_comment'] += 1
                                            else:
                                                if idx_col not in Tags:
                                                    Tags.append(idx_col)
                                                    continue





                    max_v = max(col_stat, key=col_stat.get)
                    if 'Unknown' == max_v:
                        continue
                    elif max_v == 'is_domain':
                        col_index['domain'] = idx_col
                        skip_domain = True
                        if idx_col in Tags:
                            Tags.remove(idx_col)
                    elif  max_v == 'is_owner':
                            skip_owner = True
                            col_index['owner'] = idx_col
                            if idx_col in Tags:
                                Tags.remove(idx_col)
                    elif max_v == 'is_ip':
                                skip_ip = True
                                col_index['ip'] = idx_col
                                if idx_col in Tags:
                                    Tags.remove(idx_col)
                    elif not skip_commnet:
                            col_index['comment'] = idx_col
                            if idx_col in Tags:
                                Tags.remove(idx_col)

                    if max_v == 'is_comment':
                                    if idx_col == 2:
                                            if (col_stat['is_owner']) >= 5:
                                                col_index['owner'] = idx_col
                                                if idx_col in Tags:
                                                    Tags.remove(idx_col)
                                                if settings.DEBUG:
                                                    print(
                                                        "Page: {}| Col {}| possible: [FIX] owner (d:{},i:{}, o:{},c:{},t:{})".format(
                                                            self.sheet_tag, idx_col, col_stat['is_domain'],
                                                            col_stat['is_ip'],
                                                            col_stat['is_owner'], col_stat['is_comment'], col_stat['is_tag']))
                                                continue
                                    else:
                                        skip_commnet = True


                        #print("Page: {}| Col {}| possible empty, skipped".format(self.sheet_tag, idx_col))
                    #else:
                    if DEBUG:
                        print("Page: {}| Col {}| possible: {}(d:{},i:{}, o:{},c:{},t:{})".format(self.sheet_tag, idx_col,
                                                                                                max_v,
                                                                                                col_stat['is_domain'],
                                                                                                col_stat['is_ip'],
                                                                                                col_stat['is_owner'],
                                                                                                col_stat['is_comment'],
                                                                                                col_stat['is_tag']))

        if not DEBUG:
            return self.ExtractIPInfo(domain_idx=col_index['domain'], ip_idx=col_index['ip'], owner_idx=col_index['owner'],
                                  comment_idx=col_index['comment'], stop_recurse=True, HasTags=Tags)

def UpdateCallBackStatus(taskid, name, value, event = 1):
    """Функция для создания сообщения callback на запрос статуса потоками"""
    JOB = cache.get(taskid, {})
    if event == 1:
        JOB.update({name: {'status': value}})
    else:
        JOB.update({name: {'error': value}})
    cache.set(taskid, JOB)


def make_doc(request=None, data_set={}, fileuuid='')->str:
    """Функция для генерации docx файла"""
    #request.session['docx_download_status'] = 'Создаем файл...'
    TEMPLATE_FILE = os.path.join(BASE_DIR, 'templates//ACL.docx')
    if fileuuid == '':
        fileuuid = str(uuid.uuid4())
    UpdateCallBackStatus(fileuuid, 'docx_download_status', 'Создаем docx файл')
    APP_FILE = 'static/docx/Application_' + fileuuid + '.docx'
    doc = Document(TEMPLATE_FILE)
    UpdateCallBackStatus(fileuuid, 'docx_download_status', 'Записываем изменения')

    doc.styles['Normal'].font.name = 'Verdana'
    doc.styles['Normal'].font.size = Pt(10)
    for data_inx, data in enumerate(FORM_APPLICATION_KEYS):
        row_cnt = 1
        table_tmp = doc.tables[data_inx]  # Берем таблицу по номеру
        table_rows = len(table_tmp.rows) - 1  #Берем все строки из таблицы
        if data_inx == 0:  # Для таблицы контакты, меняем правила игры
            for row_idx, row_data in enumerate(data_set[data]):
                if row_idx >= table_rows:
                    table_tmp.add_row()
                table_tmp.cell(row_idx, 1).text = row_data
        else:
            row_cnt = 0
            if data not in data_set:
                continue
            for key, value in enumerate(data_set[data], start=1):
                if row_cnt >= table_rows:
                        table_tmp.add_row()
                for cell_idx, cell_val in enumerate(value):
                    table_tmp.cell(key, cell_idx).text = cell_val

                row_cnt += 1

    if 'taskid' in request.session:
        if (request.session['taskid'] != ''):
            p = doc.add_paragraph(request.session['taskid'])
            table = doc.tables[0]
            table._element.addnext(p._p)

    #request.session['docx_download_status'] = "Сохраняем файл ..."
    UpdateCallBackStatus(fileuuid, 'docx_download_status', "Сохраняем файл {} ".format("Application_" + fileuuid))
    doc.save(os.path.join(BASE_DIR, APP_FILE))
    #request.session['docx_download_status'] = "{}".format('\\' + APP_FILE)
    #UpdateCallBackStatus(fileuuid, 'docx_download_status', "{}".format('\\' + APP_FILE))
    request.session.modified = True
    return '/' + APP_FILE


def is_valid_uuid(uuid_to_test, version=4):
    try:
        uuid_obj = uuid.UUID(uuid_to_test, version=version)
    except ValueError:
        return False
    return str(uuid_obj) == uuid_to_test


def evalute_field(record, field_spec):
    """
    Evalute a field of a record using the type of the field_spec as a guide.
    """
    if type(field_spec) is int:
        return str(record[field_spec])
    elif type(field_spec) is str:
        return str(getattr(record, field_spec))
    else:
        return str(field_spec(record))


def table(records, fields, headings=None, alignment=None, file=None):
    """
    https[:]//stackoverflow[.]com/questions/13394140/generate-markdown-tables
    Generate a Doxygen-flavor Markdown table from records.

    file -- Any object with a 'write' method that takes a single string
        parameter.
    records -- Iterable.  Rows will be generated from this.
    fields -- List of fields for each row.  Each entry may be an integer,
        string or a function.  If the entry is an integer, it is assumed to be
        an index of each record.  If the entry is a string, it is assumed to be
        a field of each record.  If the entry is a function, it is called with
        the record and its return value is taken as the value of the field.
    headings -- List of column headings.
    alignment - List of pairs alignment characters.  The first of the pair
        specifies the alignment of the header, (Doxygen won't respect this, but
        it might look good, the second specifies the alignment of the cells in
        the column.

        Possible alignment characters are:
            '<' = Left align (default for cells)
            '>' = Right align
            '^' = Center (default for column headings)
    """

    num_columns = len(fields)
    if headings:
        assert len(headings) == num_columns

    # Compute the table cell data
    columns = [[] for i in range(num_columns)]
    for record in records:
        for i, field in enumerate(fields):
            columns[i].append(evalute_field(record, field))

    # Fill out any missing alignment characters.
    extended_align = alignment if alignment != None else []
    if len(extended_align) > num_columns:
        extended_align = extended_align[0:num_columns]
    elif len(extended_align) < num_columns:
        extended_align += [('^', '<')
                           for i in range[num_columns - len(extended_align)]]

    heading_align, cell_align = [x for x in zip(*extended_align)]

    field_widths = [len(max(column, key=len)) if len(column) > 0 else 0
                    for column in columns]
    if headings:
        heading_widths = [max(len(head), 2) for head in headings]

    else:
        heading_widths = field_widths

    column_widths = [max(x) for x in zip(field_widths, heading_widths)]

    _ = ' | '.join(['{:' + a + str(w) + '}'
                    for a, w in zip(heading_align, column_widths)])
    heading_template = '| ' + _ + ' |'
    _ = ' | '.join(['{:' + a + str(w) + '}'
                    for a, w in zip(cell_align, column_widths)])

    row_template = '| ' + _ + ' |'

    _ = ' | '.join([left_rule[a] + '-' * (w - 2) + right_rule[a]
                    for a, w in zip(cell_align, column_widths)])
    ruling = '| ' + _ + ' |'

    if file is not None:
        if headings:
            file.write(heading_template.format(*headings).rstrip() + '\n')
        file.write(ruling.rstrip() + '\n')
        for row in zip(*columns):
            file.write(row_template.format(*row).rstrip() + '\n')
        file.write('\n')
        file.write('\n')


def MakeMarkDown(request, json_data, filename, fileuuid=''):
    """Функция записывает JSON как md файл"""
    #request.session['git_upload_status'] = 'Создание md-файла'
    UpdateCallBackStatus(fileuuid, 'git_upload_status', 'Создание md-файла')
    try:
        tmp = os.path.join(BASE_DIR, 'static/md/' + filename + '.md')
        file = codecs.open(tmp, "w", encoding="utf-8")
        data = json_data #json.loads(json_data)
        for key in data:
            if key == 'acl_create_info.html':
                file.write('## {}'.format(data[key][4]))
                file.write('\n')
                file.write('##### Описание доступа к ресурсам')
                file.write('\n')
                tmp = zip(contact_table, data[key])
                fields = [0, 1]
                table(records=tmp, fields=fields, headings=contact_column, alignment=[('<', '<'), ('<', '<')], file=file)

            elif key == 'acl_internal_resources.html':
                file.write('\n')
                file.write('##### Список внутренних ресурсов (СГ АльфаСтрахование)')
                file.write('\n')
                fields = [0, 1, 2]
                table(records=data[key], fields=fields, headings=standart_column,
                      alignment=[('<', '<'), ('^', '^'), ('<', '<')], file=file)


            elif key == 'acl_dmz_resources.html':
                file.write('\n')
                file.write('##### Список DMZ ресурсов (СГ АльфаСтрахование)')
                file.write('\n')
                fields = [0, 1, 2]
                table(records=data[key], fields=fields, headings=standart_column,
                      alignment=[('<', '<'), ('^', '^'), ('<', '<')], file=file)

            elif key == 'acl_external_resources.html':
                file.write('\n')
                file.write('##### Список внешних ресурсов (Internet)')
                file.write('\n')
                fields = [0, 1, 2]
                table(records=data[key], fields=fields, headings=standart_column,
                      alignment=[('<', '<'), ('^', '^'), ('<', '<')], file=file)

            elif key == 'acl_traffic.html':
                file.write('\n')
                file.write('##### Потоки трафика')
                file.write('\n')
                fields = [0, 1, 2, 3, 4, 5]
                table(records=data[key], fields=fields, headings=traffic_column,
                      alignment=[('<', '<'), ('<', '<'), ('<', '<'), ('<', '<'), ('<', '<'), ('<', '<'), ], file=file)
        if file:
            file.close()
    except Exception as e:
        logger.error('{}'.format(e))
        return False
    UpdateCallBackStatus(fileuuid, 'git_upload_status', 'ACL файл cоздан')
    #request.session['git_upload_status'] = '/static/md/' + filename + '.md'
    #request.session.modified = True
    return '/static/md/' + filename + '.md'



class GitWorker:
    def __init__(self, request, GITPRO: None, USERNAME: None, PASSWORD: None,  PATH_OF_GIT_REPO, MDFILE: None, taskid=''):
        uid = str(uuid.uuid4())
        if PATH_OF_GIT_REPO is not None:
            try:
                self.repo = git.Repo.init(PATH_OF_GIT_REPO, bare=True) #PATH_OF_GIT_REPO
            except:
                UpdateCallBackStatus(taskid, 'git_upload_status', 'Ошибка при инициализации Repo: {}'.format(PATH_OF_GIT_REPO), 0)
                return False
            if settings.DEBUG:
                logger.debug('Инициализация GIT репозитория {}'.format(PATH_OF_GIT_REPO))
        else:
            tmp = os.path.join(tempfile.gettempdir(), uid)
            try:
                self.repo = git.Repo.init(tmp, bare=True)  # uid, bare=True os.path.join(tempfile.gettempdir(), uid)
            except:
                UpdateCallBackStatus(taskid, 'git_upload_status', 'Ошибка при инициализации Repo: {}'.format(tmp), 0)
                return False
            if settings.DEBUG:
                logger.debug('Инициализация GIT репозитория {}'.format(os.path.join(tempfile.gettempdir(), uid)))

        self.request = request
        self.taskid = taskid
        #self.request.session['git_upload_status'].append({'status': "Инициализация Git проекта"})
        UpdateCallBackStatus(taskid, 'git_upload_status', 'Инициализация Git проекта')

        if PASSWORD is not None and USERNAME is not None:
             self.USERNAME = USERNAME

             if '@' in self.USERNAME:
                self.USERNAME = self.USERNAME.replace('@', '%40')
             else:
                 self.USERNAME = self.USERNAME + '%40' + 'alf'+'ast' + 'rah'+'.ru'

             if PASSWORD:
                 self.PASSWORD = PASSWORD
             if '@' in self.PASSWORD:
                 logger.warning('В пароле пользователя {} имеется запрещенный символ'.format(self.USERNAME))

             self.GITURL = GITPRO
             self.GITPRO = GITPRO.split('://')[1]
             self.GITPRO = f"https://{self.USERNAME}:{self.PASSWORD}@{self.GITPRO}"

        if settings.DEBUG:
            logger.debug('Проверка логина и пароля: {}'.format(self.GITPRO))


        if PATH_OF_GIT_REPO is not None:
            self.PATH_OF_GIT_REPO = PATH_OF_GIT_REPO
        else:
             #if settings.DEBUG:
                 #self.PATH_OF_GIT_REPO = os.path.join(os.path.abspath(os.getcwd()), str(uuid.uuid4()))
            # else:
                self.PATH_OF_GIT_REPO = os.path.join(tempfile.gettempdir(), uid)

        self.PATH_OF_GIT_REPO = os.path.join(self.PATH_OF_GIT_REPO, 'REPO')

        if not os.path.exists(self.PATH_OF_GIT_REPO):
                 os.makedirs(self.PATH_OF_GIT_REPO)
                 UpdateCallBackStatus(taskid, 'git_upload_status', "Создание временой папки")
                 #self.request.session['git_upload_status'].append({'status': "Создание временой папки: {}".format(self.PATH_OF_GIT_REPO)})
                 if settings.DEBUG:
                     logger.debug("Создание временой папки: {}".format(self.PATH_OF_GIT_REPO))
        #else:
              # os.path.join(BASE_DIR, 'upload')
        if not os.path.exists(MDFILE):
            self.MDFILE = os.path.join(os.path.abspath(os.getcwd()), MDFILE)
        else:
            self.MDFILE = MDFILE
        if settings.DEBUG:
            logger.debug("Путь к md файлу: {}".format(self.MDFILE))


    def free(self):
        for i in range(1, 3):
            self.repo.close()
            self.repo.__del__()
            if shutil.rmtree(Path(self.PATH_OF_GIT_REPO).parent, ignore_errors=True):
                break
            else:
                time.sleep(i)

    def clone(self):
        try:
            if settings.DEBUG:
                logger.debug('Копируем репозиторий: {} ->{} '.format(self.GITPRO, self.PATH_OF_GIT_REPO))
            UpdateCallBackStatus(self.taskid, 'git_upload_status', "Клонируем удаленный репозиторий")
            self.repo = self.repo.clone_from(self.GITPRO, self.PATH_OF_GIT_REPO)
        except Exception as e:
            if e.status == 128:
                #self.request.session['git_upload_status'].append({'error': "Нет доступа к GIT репозиторию"})
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Нет доступа к GIT репозиторию", 0)
            else:
                #self.request.session['git_upload_status'].append({'error': "[Ошибка] {}".format(e)})
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка при клонировании резозитория", 0)
            if settings.DEBUG:
                logger.debug('Ошибка при копировании: {}'.format(e))
            return 0

        if len(self.repo.index.entries) == 0:
           #self.request.session['git_upload_status'].append({'error': "Не удалось скачать файлы проекта, папка пустая"})
           UpdateCallBackStatus(self.taskid, 'git_upload_status', "Не удалось скачать файлы проекта, проект пустой", 0)
           if settings.DEBUG:
               logger.debug("Не удалось скачать файлы проекта, папка пустая")
           return 0
        #self.request.session['git_upload_status'].append({'status': "Скачано: {} файлов".format(len(self.repo.index.entries))})
        UpdateCallBackStatus(self.taskid, 'git_upload_status', "Скачано: {} файлов".format(len(self.repo.index.entries)))
        if settings.DEBUG:
            logger.debug("Скачано: {} файлов".format(len(self.repo.index.entries)))
        return True

    def activity(self):
        dfile = ''
        try:
            sfile = self.MDFILE
            dfile = os.path.join(self.PATH_OF_GIT_REPO, 'acl.md')
            if not copyfile(sfile, dfile):
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка при копировании md файла в проект", 0)
                #self.request.session['git_upload_status'].append({'error': "Ошибка при копировании файла в проект: {}".format(dfile)})
                if settings.DEBUG:
                    logger.debug("Ошибка при копировании файла {} в проект: {}".format(sfile, dfile))
                return 0
            UpdateCallBackStatus(self.taskid, 'git_upload_status', "Копирование md файла в проект")
            #self.request.session['git_upload_status'].append({'status': "Копирование файла в проект: {}".format(dfile)})
            if settings.DEBUG:
                logger.debug("Копирование файла в проект: {}".format(dfile))
        except Exception as e:
            #self.request.session['git_upload_status'].append({'error': "Возникла ошибка при копировании md файла в папку проекта"})
            UpdateCallBackStatus(self.taskid, 'git_upload_status', "Возникла ошибка при копировании md файла в папку проекта", 0)
            if settings.DEBUG:
                logger.debug("Возникла ошибка при копировании md файла в папку проекта: {}".format(e))
            return 0
        #finally:
        #if 'linux' in sys.platform:
            #return str(PurePosixPath(dfile)).replace('/', '//')
        return dfile #str(PurePosixPath(dfile)).replace('/', '//')

    def addindex(self, filename):
        try:
            index = self.repo.index
            index.add([filename])
            index.commit(COMMIT_MESSAGE)
            #self.request.session['git_upload_status'].append({'status': "Локальный коммит изменений"})
            UpdateCallBackStatus(self.taskid, 'git_upload_status',
                                 "Локальный коммит изменений")
            if settings.DEBUG:
                logger.debug("Локальный коммит изменений")
        except Exception as e:
                #self.request.session['git_upload_status'].append({'error': "Ошибка при коммите: {}".format(e)})
                UpdateCallBackStatus(self.taskid, 'git_upload_status',
                                     "Ошибка при локальном коммите", 0)
                if settings.DEBUG:
                    logger.debug("Ошибка при коммите: {}".format(e))
                return False
        return True

    def push(self, refspec=''):
        #remote = self.repo.create_remote('origin', self.repo.remotes.origin.url)
        if settings.DEBUG:
            logger.debug("Отправка изменений на сервер")
        try:
            if refspec == '':
                refspec = 'master:master'
            result = self.repo.remotes.origin.push(refspec=refspec)
            if result:
                #self.request.session['git_upload_status'].append({'status': "Отправка изменений на сервер"})
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Файл acl.md успешно загружен в репозиторий")
                return True
        except Exception as e:
            if e.status == 128:
                #self.request.session['git_upload_status'].append({'error': "Ошибка аутентификации для данного репозитория"})
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка аутентификации для данного репозитория", 0)
                if settings.DEBUG:
                    logger.debug("Ошибка аутентификации для данного репозитория")
            elif "src refspec master" in e.stderr:
                if refspec == 'master:master':
                    UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка при отправке в ветку master, пробуем в последний коммит: {}".format(self.repo.active_branch))
                    time.sleep(3)
                    self.push("{}:{}".format("HEAD", self.repo.active_branch))
                else:
                    UpdateCallBackStatus(self.taskid, 'git_upload_status',
                                         "Ошибка при отправке в ветку {}".format(refspec), 0)
            else:
                #self.request.session['git_upload_status'].append({'error': "Ошибка при отправки файла в проект: {}".format(e)})
                UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка при отправке коммита на сервер", 0)
                if settings.DEBUG:
                    logger.debug("Ошибка при отправки файла в проект: {}".format(e))
            return False

        finally:
            self.repo.close()
            #UpdateCallBackStatus(self.taskid, 'git_upload_status', "Ошибка при отправке коммита на сервер", 0)
            if settings.DEBUG:
                logger.debug("Очистка временной папки")
        return True


def dns_fileHandler(fname, full_buf)->int:
    """Функция парсинга DNS Aktus"""
    count = 0
    EMAIL = '(\w|\.|\_|\-)+[@](\w|\_|\-|\.)+[.]\w{2,3}'
    DEFAULT = 'Default'
    with codecs.open(fname, 'r', encoding='utf-8', errors='ignore') as f: #encoding='utf-8'
        owner = ''
        for line in f:
            try:
                #line = line.decode('cp1252').encode('utf-8')
                if len(line) > 1:
                    if line[0] not in ['@', ';', ' ', '_']:
                        if 'IN' in line:
                            if owner != '':
                                if owner not in full_buf:
                                    full_buf[owner] = []
                                full_buf[owner].append(line.strip())
                                count += 1
                            else:
                                if DEFAULT not in full_buf:
                                    full_buf[DEFAULT] = []
                                full_buf[DEFAULT].append(line.strip())
                                count += 1
                    else:
                        if ('<<' in line) or \
                             (';<' in line and not ('/' in line or '\\' in line)) or \
                             ('<' in line and re.search(EMAIL, line)) or \
                             ('OT' in line and (':' in line or '#' in line)):
                                owner = line.strip()
                        elif ('>>' in line and owner != '') or  (';<' in line and ('/' in line or '\\' in line)):
                            if owner not in full_buf:
                                full_buf[owner] = []
                            owner = ''
                        else:
                            if DEFAULT not in full_buf:
                                full_buf[DEFAULT] = []
                            if line[0] not in ['@', ';', ' ', '_']:
                                full_buf[DEFAULT].append(line.strip())
                                count += 1
            except:
                pass

    return count


def ExtractDataDns(uploaded_file_url)->int:
    """Функция записи из буфера DNS файла"""
    buff = {}
    Tags = apps.get_model('ownerlist', 'Tags')
    Iplist = apps.get_model('ownerlist', 'Iplist')
    Owners = apps.get_model('ownerlist', 'Owners')
    count = 0
    result = dns_fileHandler(uploaded_file_url, buff)
    if buff:
           for owner in buff:
               for value in buff[owner]:
                       line = value.split()
                       if len(line) < 4:
                           continue
                       elif len(line) >= 5:
                           del line[1]
                       if line[2] == 'CNAME':
                            for cname in buff:
                                for tmp in buff[cname]:
                                    s = tmp.split()
                                    if line[3] == s[0]:
                                        owner = "{} CNAME {}".format(owner, line[3])
                                        if len(s) >= 5:
                                            line[3] = s[4]
                                        else:
                                             line[3] = s[3]
                                        break
                            if not isvalidip(line[3]):
                                try:
                                    r = socket.gethostbyname(line[3])
                                    if r:
                                        owner = "{} CNAME {}".format(owner, line[3])
                                        line[3] = r
                                except:
                                    pass

                       if settings.DEBUG:
                            print(line)

                       try:
                              owner_info, created1 = Owners.objects.get_or_create(username='Макаренко А.Б')

                       except:
                               owner_info = Owners.get_default_owner()
                       try:
                           tag_info, created2 = Tags.objects.get_or_create(name='Aktur DNS')
                       except:
                            pass

                       try:
                            created3 = None
                            ip_info, created3 = Iplist.objects.get_or_create(
                                ipv4=line[3],
                                hostname=line[0],
                                owner=owner_info,
                                comment=owner,
                                )
                            if created3:
                                count += 1
                       except IntegrityError:
                            ip_info = Iplist.objects.get(ipv4=line[3])
                            ip_info.ipv4 = line[3]
                            ip_info.hostname = line[0]
                            ip_info.owner = owner_info
                            ip_info.comment = owner

                            ip_info.save()

                       except DataError as e:
                            if settings.DEBUG:
                                print("- Ошибка данных: {}".format(e))

                       if created3:
                           ip_info.tags.add(tag_info)
                           ip_info.save()


    return count


def ClearSessionMeta(request=None):
    """Функция очистки сессии при переходе на другую страницу"""
    if request:
        if 'LOCAL_STORAGE' in request.session:
            request.session['LOCAL_STORAGE'] = {}
        if 'uuid' in request.session:
            del request.session['uuid']
        if 'taskid' in request.session:
            del request.session['taskid']
        if 'GIT_URL' in request.session:
           del request.session['GIT_URL']
        if 'ACT_MAKE_GIT' in request.session:
           del request.session['ACT_MAKE_GIT']
        if 'ACT_MAKE_DOCX' in request.session:
           del request.session['ACT_MAKE_DOCX']

        if 'file_download' in request.session:
            try:
                BASE = os.path.basename(request.session['file_download'])
                if BASE:
                    BASE = os.path.join(settings.BASE_DIR, 'static//docx//' + BASE)
                    if os.path.exists(BASE):
                        os.remove(BASE)
            finally:
                del request.session['file_download']
                BASE = None

        if 'file_download_md' in request.session:
            try:
                BASE = os.path.basename(request.session['file_download_md'])
                if BASE:
                    BASE = os.path.join(settings.BASE_DIR, 'static//md//' + BASE)
                    if os.path.exists(BASE):
                        os.remove(BASE)
            finally:
                del request.session['file_download_md']
                BASE = None


def MakeTemporaryToken():
    s = "{} ACL token {}".format(random.randrange(999), datetime.datetime.now())
    return "{}".format(hashlib.md5(s.encode()).hexdigest()[:10])